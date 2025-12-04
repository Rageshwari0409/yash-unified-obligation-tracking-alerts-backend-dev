# ============================================================================
# Professional Report Synthesizer
# ============================================================================

import logging
from typing import Any, Dict, List, Literal, Optional
from datetime import datetime
import os
import base64
from io import BytesIO
import re
from litellm import completion

logger = logging.getLogger(__name__)

# --- Premium "Midnight & Gold" Palette ---
PREMIUM_PALETTE = {
    # Backgrounds
    "cover_bg": "#0F172A",      # Deepest Navy/Slate (Almost Black)
    "header_bg": "#FFFFFF",     # Pure White
    "page_bg": "#FFFFFF",
    
    # Accents
    "gold_accent": "#C5A059",   # Muted Metallic Gold
    "slate_accent": "#334155",  # Slate 700
    "brand_primary": "#1E293B", # Slate 800
    
    # Text
    "text_main": "#334155",     # Slate 700 (Softer than black)
    "text_header": "#0F172A",   # Slate 900
    "text_light": "#F1F5F9",    # Slate 100
    
    # Functional
    "table_header": "#1E293B",
    "table_row_even": "#F8FAFC", # Slate 50
    "table_border": "#E2E8F0",   # Slate 200
    "insight_bg": "#F8FAFC",     # Very light gray for boxes
    "insight_border": "#C5A059"  # Gold border for insights
}

class ReportSynthesizer:
    """
    Professional report synthesizer with premium styling, sophisticated cover page,
    and multi-format support (PDF, DOCX, MD).
    """


    def __init__(self, llm_params: Dict[str, Any], token_tracker=None):
        self.palette = PREMIUM_PALETTE
        self.llm_params = llm_params
        self.token_tracker = token_tracker
    
    async def synthesize_report(
        self,
        document_context: str,
        format: Literal["pdf", "docx", "md"] = "pdf"
    ) -> Dict[str, Any]:
        """
        Generate comprehensive report with premium professional styling.
        """
        logger.info(f"Synthesizing professional report in {format} format")
        
        # 1. Generate narrative report text using LLM
        narrative = await self._generate_narrative(document_context)
        
        # 2. Generate report file based on requested format
        if format == "pdf":
            report_content = await self._generate_pdf_report(narrative)
        elif format == "docx":
            report_content = await self._generate_docx_report(narrative)
        else:
            report_content = await self._generate_markdown_report(narrative)
        
        return {
            "content": report_content,
            "format": format,
            "narrative": narrative,
            "metadata": {
                "generated_at": datetime.utcnow().isoformat(),
                "format": format,
                "theme": "midnight_gold_premium"
            }
        }
    
    async def _generate_narrative(
        self,
        document_context: str
    ) -> str:
        """Generate contract analysis narrative with obligation focus."""

        prompt = f"""You are a senior contract analyst preparing an executive summary report. Create a professional Contract Analysis Report (4-5 pages) with actionable insights.

## DOCUMENT CONTEXT:
{document_context}

## CRITICAL RULES:
1. ONLY include information explicitly found in the document
2. NO explanations or assumptions - just extracted data
3. If information is not found, write "Not specified"
4. Keep report concise: 4-5 pages maximum
5. DO NOT use tables - use bullet point format instead

## REPORT STRUCTURE:

### 1. EXECUTIVE SUMMARY (1 paragraph)
- Contract type and primary purpose
- Key parties and their roles
- Contract value and duration
- Critical obligations and deadlines

### 2. CONTRACT OVERVIEW
- **Contract Type:** [Extract from document]
- **Parties Involved:** [List all parties with roles]
- **Effective Date:** [Extract from document]
- **Expiration Date:** [Extract from document]
- **Contract Value:** [Extract from document]
- **Governing Law:** [Extract from document]

### 3. KEY TERMS & CONDITIONS
- **Payment Terms:** [Extract exact terms]
- **Delivery/Performance:** [Extract exact terms]
- **Warranties & Representations:** [Extract exact terms]
- **Intellectual Property:** [Extract exact terms]
- **Confidentiality:** [Extract exact terms]
- **Termination:** [Extract exact terms]

### 4. OBLIGATIONS SUMMARY
List each obligation in this format:

**Obligation 1:** [Description]
- **Responsible Party:** [Party name]
- **Due Date:** [Date or "Not specified"]
- **Priority:** [High/Medium/Low]
- **Status:** [Pending/Complete]

**Obligation 2:** [Description]
- **Responsible Party:** [Party name]
- **Due Date:** [Date or "Not specified"]
- **Priority:** [High/Medium/Low]
- **Status:** [Pending/Complete]

[Continue for all obligations found]

### 5. RISK ASSESSMENT
**HIGH RISK:**
- [List only if found in document]

**MEDIUM RISK:**
- [List only if found in document]

**LOW RISK:**
- [List only if found in document]

### 6. COMPLIANCE REQUIREMENTS
- [Extract from document only]

### 7. CRITICAL DEADLINES
List in chronological order:
- **[Date]:** [Obligation] - [Responsible Party]
- **[Date]:** [Obligation] - [Responsible Party]

### 8. ACTION ITEMS & RECOMMENDATIONS
**Immediate Priority:**
- [Action] - Owner: [Party] - Deadline: [Date]

**Short-term:**
- [Action] - Owner: [Party] - Deadline: [Date]

**Ongoing:**
- [Action] - Owner: [Party]

## FORMATTING:
- Use `##` for main headers, `###` for sub-headers
- Use `**bold**` for labels and important terms
- Use bullet points (`-`) for lists
- Dates in YYYY-MM-DD format
- NO tables - use structured bullet points only
- Be concise - no filler text

## OUTPUT:
Generate the complete report in Markdown format. No preamble.

---

# CONTRACT ANALYSIS REPORT
"""
        
        try:
            response = completion(
                **self.llm_params,
                messages=[
                    {
                        "content": prompt,
                        "role": "user"
                    }
                ],
            )
            if self.token_tracker:
                self.token_tracker.track_response(response)
            narrative = response.choices[0].message.content
            logger.info("Generated enhanced report narrative")
            return narrative

        except Exception as e:
            logger.error(f"Narrative generation failed: {e}")
            return "Error generating report narrative."

    # =========================================================================
    # PDF GENERATION HELPERS
    # =========================================================================

    def _draw_geometric_pattern(self, canvas, x, y, width, height, color, alpha=0.05):
        """Draws a subtle premium geometric pattern (dots + tech lines) on the cover."""
        canvas.saveState()
        canvas.setFillColor(color)
        canvas.setStrokeColor(color)
        canvas.setFillAlpha(alpha)
        canvas.setStrokeAlpha(alpha)
        
        # 1. Dot Matrix Pattern
        step = 40
        for i in range(0, int(width), step):
            for j in range(0, int(height), step):
                if (i + j) % (step * 2) == 0:
                    canvas.circle(x + i, y + j, 1.5, fill=1)
        
        # 2. Subtle Geometric Angles (Tech feel)
        canvas.setLineWidth(1)
        path = canvas.beginPath()
        path.moveTo(width, 0)
        path.lineTo(width, height * 0.4)
        path.lineTo(0, 0)
        path.close()
        canvas.drawPath(path, stroke=0, fill=1)
        
        canvas.restoreState()

    def _draw_cover(self, canvas, doc):
        """Render the darker, premium cover page."""
        from reportlab.lib.colors import HexColor
        from reportlab.lib.pagesizes import A4
        
        canvas.saveState()
        width, height = A4
        
        # Colors
        bg_color = HexColor(self.palette["cover_bg"])
        gold_color = HexColor(self.palette["gold_accent"])
        white_color = HexColor("#FFFFFF")
        slate_text = HexColor("#94A3B8")
        
        # 1. Full Dark Background
        canvas.setFillColor(bg_color)
        canvas.rect(0, 0, width, height, stroke=0, fill=1)
        
        # 2. Texture (Subtle white dots on dark bg)
        self._draw_geometric_pattern(canvas, 0, 0, width, height, white_color, alpha=0.03)
        
        # 3. Gold Accent Bar (Left side)
        canvas.setFillColor(gold_color)
        canvas.rect(0, 0, 8, height, stroke=0, fill=1)
        
        # 4. Header / Brand Area
        canvas.setFont("Helvetica-Bold", 10)
        canvas.setFillColor(gold_color)
        canvas.drawString(40, height - 50, "CONTRACT OBLIGATION ANALYSIS")
        
        # 5. Footer Info (Bottom)
        # Line separator
        canvas.setStrokeColor(HexColor("#334155"))
        canvas.setLineWidth(1)
        canvas.line(40, 60, width - 40, 60)
        
        # Text
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(slate_text)
        canvas.drawString(40, 40, "CONFIDENTIAL & PROPRIETARY")
        canvas.drawRightString(width - 40, 40, f"REF: {datetime.utcnow().strftime('%Y-%m-%d')}")

        canvas.restoreState()

    def _draw_content_header(self, canvas, doc):
        """Render the minimalist clean header for content pages."""
        from reportlab.lib.colors import HexColor
        from reportlab.lib.pagesizes import A4
        
        canvas.saveState()
        width, height = A4
        
        header_bg = HexColor(self.palette["header_bg"])
        primary = HexColor(self.palette["brand_primary"])
        accent = HexColor(self.palette["gold_accent"])
        
        # White Header Strip
        header_height = 50
        canvas.setFillColor(header_bg)
        canvas.rect(0, height - header_height, width, header_height, stroke=0, fill=1)
        
        # Accent Line (Bottom of header)
        canvas.setStrokeColor(accent)
        canvas.setLineWidth(1)
        canvas.line(0, height - header_height, width, height - header_height)
        
        # Text
        canvas.setFont("Helvetica-Bold", 9)
        canvas.setFillColor(primary)
        canvas.drawString(40, height - 32, "CONTRACT ANALYSIS REPORT")
        
        # Page Number
        page_num = f"{doc.page}"
        canvas.setFont("Helvetica", 9)
        canvas.drawRightString(width - 40, height - 32, f"Page {page_num}")
        
        canvas.restoreState()

    def _sanitize_text(self, text):
        """Remove broken/unpaired HTML tags that cause ReportLab parsing errors."""
        # Remove all HTML-like tags except allowed ones
        allowed_tags = ['b', 'i', 'u', 'br', 'para', 'font', 'super', 'sub']

        # First, escape < and > that might cause issues
        text = text.replace('</i>', ' ').replace('<i>', ' ')
        text = text.replace('</b>', ' ').replace('<b>', ' ')

        # Now re-apply proper bold/italic from markdown
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)

        # Remove any stray < or > not part of valid tags
        text = text.replace('<', '').replace('>', '')

        return text

    def _parse_markdown_table(self, table_lines, styles):
        """Parse markdown table and return a ReportLab Table flowable."""
        from reportlab.platypus import Table, TableStyle, Paragraph
        from reportlab.lib.colors import HexColor
        from reportlab.lib import colors

        if len(table_lines) < 2:
            return None

        table_data = []
        for i, line in enumerate(table_lines):
            # Skip separator line (|---|---|)
            if re.match(r'^\|[\s:-]+\|$', line.replace(' ', '')):
                continue

            # Parse cells
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                # Format cell content
                formatted_cells = []
                for cell in cells:
                    cell_fmt = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', cell)
                    cell_fmt = self._sanitize_text(cell_fmt)
                    # Wrap in Paragraph for proper text wrapping
                    formatted_cells.append(Paragraph(cell_fmt, styles['Body_Clean']))
                table_data.append(formatted_cells)

        if not table_data:
            return None

        # Calculate column widths based on content
        num_cols = len(table_data[0]) if table_data else 0
        col_width = 500 / max(num_cols, 1)  # Distribute evenly

        table = Table(table_data, colWidths=[col_width] * num_cols)

        # Style the table
        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor(self.palette["table_header"])),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, HexColor(self.palette["table_border"])),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, HexColor(self.palette["table_row_even"])]),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ])
        table.setStyle(table_style)

        return table

    def _parse_markdown_to_flowables(self, text, styles):
        """
        Robust Markdown parser that handles headers, bold/italic, tables,
        and prevents bullet point bunching by properly managing ListFlowables.
        """
        from reportlab.platypus import Paragraph, ListFlowable, ListItem, HRFlowable, Spacer

        flowables = []
        lines = text.split('\n')

        # Buffers
        list_buffer = []
        table_buffer = []
        in_table = False

        for line in lines:
            line_stripped = line.strip()

            # Detect table lines (starts with |)
            is_table_line = line_stripped.startswith('|') and line_stripped.endswith('|')

            # Handle table accumulation
            if is_table_line:
                # Flush list buffer first
                if list_buffer:
                    flowables.append(ListFlowable(
                        list_buffer, bulletType='bullet', start='circle',
                        leftIndent=20, spaceAfter=12
                    ))
                    list_buffer = []

                table_buffer.append(line_stripped)
                in_table = True
                continue
            elif in_table and table_buffer:
                # End of table - parse and add it
                table_flowable = self._parse_markdown_table(table_buffer, styles)
                if table_flowable:
                    flowables.append(Spacer(1, 10))
                    flowables.append(table_flowable)
                    flowables.append(Spacer(1, 10))
                table_buffer = []
                in_table = False

            if not line_stripped:
                # Empty line: flush list buffer
                if list_buffer:
                    flowables.append(ListFlowable(
                        list_buffer, bulletType='bullet', start='circle',
                        leftIndent=20, spaceAfter=12
                    ))
                    list_buffer = []
                continue

            # Basic formatting
            line_fmt = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line_stripped)
            line_fmt = re.sub(r'\*(.+?)\*', r'<i>\1</i>', line_fmt)
            line_fmt = self._sanitize_text(line_fmt)

            # Headers
            if line_stripped.startswith('# '):
                flowables.append(Paragraph(line_fmt[2:], styles['H1_Premium']))
            elif line_stripped.startswith('## '):
                flowables.append(Paragraph(line_fmt[3:], styles['H1_Premium']))
            elif line_stripped.startswith('### '):
                flowables.append(Paragraph(line_fmt[4:], styles['H2_Premium']))
            elif line_stripped.startswith('#### '):
                flowables.append(Paragraph(line_fmt[5:], styles['H2_Premium']))
            elif line_stripped.startswith('##### '):
                flowables.append(Paragraph(line_fmt[6:], styles['H2_Premium']))
            elif line_stripped.startswith('###### '):
                flowables.append(Paragraph(line_fmt[7:], styles['H2_Premium']))
            elif line_stripped.startswith('---'):
                flowables.append(HRFlowable(
                    width="80%", thickness=1, lineCap='round',
                    color=self.palette['slate_accent'], spaceBefore=1, spaceAfter=1,
                    hAlign='CENTER', vAlign='BOTTOM', dash=None
                ))

            # List Items
            elif line_stripped.startswith('* ') or line_stripped.startswith('- '):
                clean_text = line_fmt[2:]
                list_buffer.append(ListItem(Paragraph(clean_text, styles['Body_Clean'])))

            # Standard Paragraphs
            else:
                if list_buffer:
                    flowables.append(ListFlowable(
                        list_buffer, bulletType='bullet', start='circle',
                        leftIndent=20, spaceAfter=12
                    ))
                    list_buffer = []
                flowables.append(Paragraph(line_fmt, styles['Body_Clean']))

        # Final flush
        if list_buffer:
            flowables.append(ListFlowable(
                list_buffer, bulletType='bullet', start='circle',
                leftIndent=20, spaceAfter=12
            ))
        if table_buffer:
            table_flowable = self._parse_markdown_table(table_buffer, styles)
            if table_flowable:
                flowables.append(Spacer(1, 10))
                flowables.append(table_flowable)

        return flowables

    async def _generate_pdf_report(
        self,
        narrative: str
    ) -> bytes:
        """Generate Premium PDF Report"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.colors import white, HexColor
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import (
                BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, 
                PageBreak, Table, TableStyle, NextPageTemplate, Image, KeepTogether
            )
            from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
            
            buffer = BytesIO()
            
            # 1. Document Template
            doc = BaseDocTemplate(
                buffer,
                pagesize=A4,
                topMargin=1*inch, 
                bottomMargin=1*inch, 
                leftMargin=0.75*inch, 
                rightMargin=0.75*inch
            )
            
            # 2. Frames
            full_frame = Frame(
                doc.leftMargin, 
                doc.bottomMargin, 
                doc.width, 
                doc.height - 0.5*inch, 
                id='normal'
            )
            
            # Cover Frame (Offset for the gold bar)
            cover_frame = Frame(
                doc.leftMargin + 0.5*inch, 
                doc.bottomMargin, 
                doc.width - 0.5*inch, 
                doc.height, 
                id='cover'
            )

            doc.addPageTemplates([
                PageTemplate(id='CoverTemplate', frames=[cover_frame], onPage=self._draw_cover),
                PageTemplate(id='ContentTemplate', frames=[full_frame], onPage=self._draw_content_header)
            ])

            # 3. Styles
            styles = getSampleStyleSheet()
            
            # Colors
            c_primary = HexColor(self.palette["brand_primary"])
            c_gold = HexColor(self.palette["gold_accent"])
            c_text = HexColor(self.palette["text_main"])
            
            # Cover Styles
            styles.add(ParagraphStyle(
                name='CoverTitle',
                fontName='Helvetica-Bold',
                fontSize=44,
                leading=50,
                textColor=white,
                spaceAfter=25
            ))
            
            styles.add(ParagraphStyle(
                name='CoverSubtitle',
                fontName='Helvetica',
                fontSize=18,
                leading=24,
                textColor=c_gold,
                spaceAfter=50
            ))
            
            # Content Styles
            styles.add(ParagraphStyle(
                name='H1_Premium',
                parent=styles['Heading1'],
                fontName='Helvetica-Bold',
                fontSize=20,
                leading=26,
                textColor=c_primary,
                spaceBefore=24,
                spaceAfter=12,
                borderPadding=6,
                borderWidth=0,
                borderBottomWidth=1,
                borderColor=c_gold # Gold underline
            ))
            
            styles.add(ParagraphStyle(
                name='H2_Premium',
                parent=styles['Heading2'],
                fontName='Helvetica-Bold',
                fontSize=14,
                leading=18,
                textColor=c_primary,
                spaceBefore=18,
                spaceAfter=8
            ))
            
            styles.add(ParagraphStyle(
                name='Body_Clean',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=10.5,
                leading=16,
                textColor=c_text,
                alignment=TA_JUSTIFY,
                spaceAfter=10
            ))
            
            styles.add(ParagraphStyle(
                name='InsightCallout',
                parent=styles['Normal'],
                fontName='Helvetica-Oblique',
                fontSize=10,
                leading=15,
                textColor=c_primary,
                backColor=HexColor(self.palette["insight_bg"]),
                borderPadding=12,
                borderWidth=1,
                borderColor=HexColor(self.palette["table_border"]),
                spaceBefore=10,
                spaceAfter=15
            ))
            
            styles.add(ParagraphStyle(
                name='ChartCaption',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=9,
                textColor=HexColor("#64748B"),
                alignment=TA_CENTER,
                spaceBefore=5
            ))

            # 4. Construct Story
            story = []

            # --- COVER PAGE ---
            story.append(Spacer(1, 1.5*inch))
            story.append(Paragraph("CONTRACT<br/>ANALYSIS<br/>REPORT", styles['CoverTitle']))
            story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", styles['CoverSubtitle']))

            # Contact/Meta block
            contact_style = ParagraphStyle('Contact', parent=styles['Body_Clean'], textColor=white, fontSize=11)
            story.append(Spacer(1, 1*inch))
            story.append(Paragraph("<b>DOCUMENT TYPE:</b><br/>Contract Obligation Analysis", contact_style))
            story.append(Spacer(1, 15))
            story.append(Paragraph("<b>GENERATED BY:</b><br/>Obligation Alerts System", contact_style))
            
            story.append(NextPageTemplate('ContentTemplate'))
            story.append(PageBreak())
            
            # --- NARRATIVE SECTION ---
            story.extend(self._parse_markdown_to_flowables(narrative, styles))
            story.append(PageBreak())

            doc.build(story)
            pdf_content = buffer.getvalue()
            buffer.close()
            
            logger.info("Professional PDF report generated successfully")
            return pdf_content
            
        except ImportError:
            logger.error("ReportLab not installed. Install with: pip install reportlab")
            raise
        except Exception as e:
            logger.error(f"PDF generation failed: {e}", exc_info=True)
            raise

    # =========================================================================
    # DOCX GENERATION
    # =========================================================================

    async def _generate_docx_report(
        self,
        narrative: str
    ) -> bytes:
        """Generate DOCX report with consistent styling."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Convert hex to RGB for docx
            def h2r(h): 
                h = h.lstrip('#')
                return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

            c_primary = h2r(self.palette["brand_primary"])
            c_gold = h2r(self.palette["gold_accent"])
            c_slate = h2r(self.palette["text_main"])
            
            # Title
            title = doc.add_heading('Analytical Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.color.rgb = c_primary
            title.runs[0].font.size = Pt(28)
            
            subtitle = doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%B %d, %Y')}")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.color.rgb = c_gold
            
            doc.add_paragraph()
            
            # Narrative
            for section in narrative.split('\n\n'):
                section = section.strip()
                if not section: continue
                
                if section.startswith('##'):
                    h = doc.add_heading(section.replace('#', '').strip(), level=1)
                    h.runs[0].font.color.rgb = c_primary
                elif section.startswith('###'):
                    h = doc.add_heading(section.replace('#', '').strip(), level=2)
                    h.runs[0].font.color.rgb = c_slate
                elif section.startswith('*') or section.startswith('-'):
                    for line in section.split('\n'):
                        doc.add_paragraph(line.strip()[1:].strip(), style='List Bullet')
                else:
                    doc.add_paragraph(section)

            buffer = BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
            
        except ImportError:
            logger.error("python-docx not installed")
            return b""
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            return b""

    # =========================================================================
    # MARKDOWN GENERATION
    # =========================================================================

    async def _generate_markdown_report(
        self,
        narrative: str
    ) -> str:
        """Generate enhanced Markdown report."""
        md_content = f"""# Analytical Report

**Generated:** {datetime.utcnow().strftime('%B %d, %Y')}

---

{narrative}
"""
        return md_content
    

if __name__ == "__main__":
    # this is just an example 
    import asyncio
    async def report ():
        
        report = ReportSynthesizer(
            llm_params={
                "model": "gemini/gemini-2.5-flash-lite",
                "api_key": os.getenv("GOOGLE_API_KEY")
            }
        )
        data = await report.synthesize_report("Generate types of heading of markdown", format="pdf")

        content = data['content']

        with open('report.pdf', 'wb') as f:
            f.write(content)
    
    asyncio.run(report())
